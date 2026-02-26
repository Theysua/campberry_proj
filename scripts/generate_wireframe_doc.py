from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_box(doc, text, bg_color="F5F5F5", width=None):
    table = doc.add_table(rows=1, cols=1)
    if width:
        table.autofit = False
        table.columns[0].width = Inches(width)
    
    cell = table.cell(0, 0)
    cell.text = text
    
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), bg_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        borders.append(border)
    cell._tc.get_or_add_tcPr().append(borders)
    doc.add_paragraph()

doc = Document()
doc.add_heading('Campberry Wireframes', 0)

def add_header(doc):
    add_box(doc, "[❄ Campberry (beta)]    |    🔍 Find    📋 Lists    📌 My Lists    |    [Sign In]", "E0E0E0")


## PAGE 1 - Homepage
doc.add_heading('Page 1 — Homepage ( / )', level=1)
add_header(doc)

doc.add_heading('Hero Section', level=2)
add_box(doc, "Extracurriculars & Enrichment\nTop Opportunities Curated by College Experts\n\n[ 🔍 Search by interest, name... ] [ Search ]", "D6E8FA")

doc.add_heading('Top Searches', level=2)
add_box(doc, "[Browse All] [Summer] [School Year] [Competition] [Free] [Engineering] [Leadership] [Coding] [Writing] [STEM] [Business] [Sustainability]", "FFFFFF")

doc.add_heading('Featured Lists (Carousel)', level=2)
add_box(doc, "⭐ Featured Lists                     See All Lists →\n\n[Card: School Counseling Group's Favorite Programs]\n[Card: Top STEM Programs for Rising Juniors]\n[Card: Best Free Summer Programs 2026]\n[ ← → Arrow Controls ]", "FFFFFF")

doc.add_heading('Campberry Ratings', level=2)
add_box(doc, "Experts' Choice [NEW]\nMOST / HIGHLY\n\nFinancial Accessibility [NEW]\nA+ / A / A- / B+ / B- / C+", "FFFFFF")

doc.add_heading('Trust Promises', level=2)
add_box(doc, "Trust matters. These are our promises to you:\n1. 100% Data Privacy\n2. Unbiased Algorithms\n3. Free, Forever", "EFF6FF")

doc.add_heading('Statistics', level=2)
add_box(doc, "Search 2,100+ Opportunities\nSave Time and Don't Miss Out\n\nCampberry is built by teens and educators, for teens and educators. [Meet the community]\n\n[📝 Create a List] [⭐ Leave a Review] [🚩 Flag Incorrect Info]", "FFFFFF")

doc.add_heading('Footer', level=2)
add_box(doc, "❄ Campberry\nAbout | Legal\nSocial | Contact Us\n© Campberry 2026", "EFF6FF")


## PAGE 2 - Search Results
doc.add_page_break()
doc.add_heading('Page 2 — Search Results ( /search/results )', level=1)
add_header(doc)

add_box(doc, "[ ← ] [ 🔍 Search by interest, name... ] [ 🔍 ]", "F5F5F5")

p = doc.add_paragraph()
r = p.add_run("Layout: 2 Columns (Left: Filters, Right: Results List)")
r.bold = True

add_box(doc, """[ LEFT PANEL: Filters ]
- Experts' Choice
- Type (Program/Competition)
- Location (Radius / Online)
- Season (Summer/Fall/Spring)
- Current Grade (9/10/11/12)
- Interests (STEM, Research, etc.)
- Financial Accessibility (A+ to C+)
""", "FFFFFF", 3)

add_box(doc, """[ RIGHT PANEL: Results ]
1-100 of 2100+ results    [ Sort: Relevancy ▼ ]

[Card] Clark Scholars Program (Texas Tech)
       Tags: Research, STEM, Science
       Deadline: in 1 day

[Card] MIT PRIMES (MIT)
       Tags: Math, Research
       Deadline: Dec 1, 2026

[Contact Our Experts For Free Recommendations]

[ ... More Results ... ]
[ Pagination: 1 2 3 ... 22 ]
""", "FFFFFF", 6)


## PAGE 3 - Program Detail
doc.add_page_break()
doc.add_heading('Page 3 — Program Detail ( /learning-opportunities/{uuid} )', level=1)

add_box(doc, "[ ← Back ]               [ ↗ Share ] [ ☆ Save ]", "E0E0E0")
add_box(doc, "[ YouTube Video Embed ]\n[ Thumbnail 1 ] [ Thumbnail 2 ]", "F5F5F5")
add_box(doc, """[Logo] Clark Scholars Program (Texas Tech University)
Tags: Research, STEM, Science
Summer: Jun 17 - Aug 2 | Lubbock, TX
Deadline: in 1 day | View Costs
""", "FFFFFF")
doc.add_paragraph("--- About ---")
doc.add_paragraph("--- Accordions (Sessions, Applications, Eligibility) ---")
doc.add_paragraph("[🌐 View Website]")
add_box(doc, """Ratings
- Highly Selective: Yes
- Financial Accessibility: A+
""", "FFFFFF")


## PAGE 4 - Lists Overview
doc.add_page_break()
doc.add_heading('Page 4 — Lists Overview ( /lists )', level=1)
add_header(doc)

doc.add_heading('Featured Lists (Grid)', level=2)
add_box(doc, """[Card] School Counseling Group's Favorite Programs
[Card] Top STEM Programs for Rising Juniors
[Card] Best Free Summer Programs 2026
[Card] Wharton Pre-College Programs
[Card] Leadership Programs for High Schoolers
[Card] Community Service Opportunities
""", "FFFFFF")


## PAGE 5 - List Detail
doc.add_page_break()
doc.add_heading('Page 5 — List Detail ( /lists/{uuid} )', level=1)
add_box(doc, "[ ← Back ]               [ ↗ Share ] [ Save List ]", "E0E0E0")
add_box(doc, """School Counseling Group's Favorite Programs
From Admin | Updated Dec 22, 2025
Multi-paragraph description of the list...
""", "FFFFFF")

doc.add_heading('Opportunities List', level=2)
add_box(doc, """[Card] Clark Scholars Program
[Author Commentary]: "This is one of the most prestigious research programs..."

--- Inserted Text Paragraph ---

[Card] Wharton Pre-Baccalaureate Program
[Author Commentary]: "An excellent introduction to business..."
""", "FFFFFF")


## PAGE 6 - Auth
doc.add_page_break()
doc.add_heading('Page 6 — Authentication ( /auth )', level=1)
add_box(doc, """[ Sign In ]                     [ Sign Up ]
Continue with Google            Continue with Google
----- or -----                  ----- or -----
Email                           Email
Password                        Password
[ SIGN IN ]                     [ SIGN UP ]
""", "EFF6FF")


## PAGE 7 - Mission
doc.add_page_break()
doc.add_heading('Page 7 — Mission ( /mission )', level=1)
add_header(doc)

doc.add_heading('Hero', level=2)
add_box(doc, "Our mission is to make the education sector radically more open and navigable for all.", "1E3A5F")

doc.add_heading('Promises', level=2)
add_box(doc, """1. 100% Data Privacy (We never sell your data)
2. Unbiased Algorithm (No pay-to-rank)
3. Free, Forever (All core access is totally free)
""", "EFF6FF")


doc.save('Campberry_Wireframe.docx')
print("Wireframe saved as Campberry_Wireframe.docx")
